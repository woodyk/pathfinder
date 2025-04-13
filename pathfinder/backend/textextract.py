#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: textextract.py
# Author: Wadih Khairallah
# Description: 
# Created: 2024-12-01 12:12:08
# Modified: 2025-04-05 00:22:46

import json
import math
import re
import os
import subprocess
import magic
import hashlib
import pytesseract
import requests 
import pandas as pd
import speech_recognition as sr
import fitz

from bs4 import BeautifulSoup
from collections import Counter
from docx import Document
from datetime import datetime
from mss import mss
from urllib.parse import urlparse
from io import StringIO
from PIL import Image
from pydub import AudioSegment
from rich.console import Console

console = Console()
print = console.print
log = console.log

def clean_path(path):
    path = os.path.expanduser(path)
    path = os.path.abspath(path)
    
    if os.path.isfile(path) or os.path.isdir(path):
        return path
    else:
        return False

def get_screenshot():
    ''' Take screenshot and return text object for all text found in the image '''
    # Screenshot storage path
    path = r'/tmp/sym_screenshot.png'

    with mss() as sct:
        monitor = {"top": 0, "left": 0, "width": 0, "height": 0}
        
        for mon in sct.monitors:
            # get furthest left point
            monitor["left"] = min(mon["left"], monitor["left"])
            # get highest point
            monitor["top"] = min(mon["top"], monitor["top"])
            # get furthest right point
            monitor["width"] = max(mon["width"]+mon["left"]-monitor["left"], monitor["width"])
            # get lowest point
            monitor["height"] = max(mon["height"]+mon["top"]-monitor["top"], monitor["height"])
        
        screenshot = sct.grab(monitor)

    img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
    img_gray = img.convert("L")
    img_gray.save(path)

    return path

def extract_exif(file_path):
    exif_data = None
    try:
        result = subprocess.run(['exiftool', '-j', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode())[0]

    except Exception as e:
        print(f"Exiftool failed: {e}")

    return exif_data

def text_from_url(url):
    """
    Fetch and extract text from a given URL.

    Args:
        url (str): The website URL.

    Returns:
        str: Extracted plain text from the web page.
    """
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Extract only visible text (excluding script, style, and metadata)
        for tag in soup(["script", "style", "noscript", "iframe", "header", "footer", "meta", "link"]):
            tag.decompose()

        text = soup.get_text(separator=" ")
        return text.strip()

    except requests.RequestException as e:
        print(f"Error fetching URL: {url} - {e}")
        return None

def extract_text(file_path):
    """
    Extracts text content from a file based on its MIME type.

    Processes various file types (text, Excel, PDF, Word, images,
    audio) and returns their extracted text as a UTF-8 string. Uses
    helper methods for non-text files (e.g., PDF to markdown,
    audio transcription). Logs issues and returns None if no
    content is found or an error occurs.

    Args:
        file_path (str): Path to the file, cleaned via `clean_path`.

    Returns:
        str or None: Extracted text if successful, else None.

    Raises:
        Exception: Caught internally, logged with `log` function.

    Supported MIME Types:
        - Text: 'text/*', 'application/json', 'application/xml',
          'application/x-yaml', 'text/markdown'
        - Excel: 'application/vnd.ms-excel',
          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        - PDF: 'application/pdf'
        - Word: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        - Image: 'image/*'
        - Audio: 'audio/*'

    Examples:
        >>> extract_text("example.txt")
        'Hello, world!'
        >>> extract_text("example.pdf")
        '# Page 1\\nContent here...'
        >>> extract_text("invalid_file.txt")
        None  # Logs "Error reading invalid_file.txt: [error]"
    """

    file_path = clean_path(file_path)
    if file_path is False: 
        print(f"No such file: {file_path}")
        return f"No such file: {file_path}"

    print(f"[cyan]Extracting text from:[/cyan] {file_path}")

    file_path = clean_path(file_path)
    mime_type = magic.from_file(file_path, mime=True)
    try:
        content = "" 
        if mime_type.startswith('text/') or mime_type in ['application/json', 'application/xml', 'application/x-yaml', 'text/markdown']:
            with open(file_path, 'r') as f:
                content = f.read()

        elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            content = text_from_excel(file_path)

        elif mime_type == 'application/pdf':
            content = text_from_pdf(file_path)

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = text_from_word(file_path)

        elif mime_type.startswith('image/'):
            content = text_from_image(file_path)

        elif mime_type.startswith('audio/'):
            content = text_from_audio(file_path)
        
        else:
            content = text_from_other(file_path)

        if len(content) > 0:
            content = content.encode('utf-8').decode('utf-8', errors='ignore')
            return content

        else:
            print(f"No content found for file: {file_path}")
            return None

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def text_from_audio(audio_file):
    text = ""

    def audio_to_wav(file_path):
        # Extract the file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lstrip('.')

        # Use pydub to convert to WAV
        audio = AudioSegment.from_file(file_path, format=ext)
        wav_file_path = file_path.replace(ext, 'wav')
        audio.export(wav_file_path, format='wav')

        return wav_file_path

    recognizer = sr.Recognizer()
    _, ext = os.path.splitext(audio_file)
    # Convert the file to WAV if necessary
    if ext.lower() not in ['.wav', '.wave']:
        audio_file = audio_to_wav(audio_file)
    try:
        with sr.AudioFile(audio_file) as source:
            audio = recognizer.record(source)

        text = recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        print("Google Speech Recognition could not understand audio")
        return None
    except sr.RequestError as e:
        print(f"Could not request results from Google Speech Recognition service; {e}")
        return None

    return text

def downloadImage(url):
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join('/tmp/', filename)

        response = requests.get(url, stream=True)
        response.raise_for_status()

        with open(save_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)

        return clean_path(save_path)
    else:
        print(f"Unable to pull image from {url}")
        return None

def is_image(file_path_or_url):
    try:
        mime = magic.from_file(file_path_or_url, mime=True)
        if mime.startswith("image/"):
            return true
    except Exception as e:
        return False

def text_from_pdf(pdf_path):
    """
    Extracts plain text from a PDF using PyMuPDF (fitz),
    including metadata and OCR for images.
    """
    plain_text = ""

    try:
        doc = fitz.open(pdf_path)

        # Extract metadata
        metadata = doc.metadata
        if metadata:
            for key, value in metadata.items():
                plain_text += f"{key}: {value}\n"
            plain_text += "\n"

        # Iterate through pages
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            plain_text += f"\n\n--- Page {page_num + 1} ---\n\n"

            # Extract text
            text = page.get_text()
            plain_text += text if text.strip() else "[No text found on this page]\n"

            # Extract and OCR images
            image_list = page.get_images(full=True)
            if image_list:
                for image_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_path = f"/tmp/page-{page_num + 1}-image-{image_index + 1}.png"

                    # Save image
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)

                    # Perform OCR on the image
                    image_text = text_from_image(image_path)
                    plain_text += f"\n[Extracted Text from Image {image_index + 1}]\n{image_text}\n"
            else:
                plain_text += "\n[No images found]\n"

        doc.close()

    except Exception as e:
        print(f"Error processing PDF with PyMuPDF: {pdf_path}\n{e}")
        return None

    return plain_text

def text_from_word(file_path):
    """
    Extracts plain text from a Word (.docx) file, including text, tables, and images with OCR.
    """
    file_path = clean_path(file_path)
    doc = Document(file_path)
    plain_text = ""

    # Extract text from paragraphs
    try:
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                plain_text += paragraph.text.strip() + "\n\n"
    except Exception as e:
        print(f"Error extracting text from Word file: {file_path}\n{e}")
        return None

    # Extract text from tables
    try:
        for table_num, table in enumerate(doc.tables, start=1):
            plain_text += f"\n[Table {table_num}]\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                plain_text += "\t".join(cells) + "\n"
    except Exception as e:
        print(f"Error extracting tables from Word file: {file_path}\n{e}")
        return None

    # Extract and process images
    try:
        image_num = 0
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_num += 1
                image_data = doc.part.rels[rel].target_part.blob  # Extract image data
                image_path = f"/tmp/word_image_{image_num}.png"
                
                # Save image
                with open(image_path, "wb") as img_file:
                    img_file.write(image_data)

                # Perform OCR on the extracted image
                image_text = text_from_image(image_path)
                plain_text += f"\n[Extracted Text from Image {image_num}]\n{image_text}\n"
    except Exception as e:
        print(f"Error extracting images from Word file: {file_path}\n{e}")
        return None

    return plain_text

def text_from_excel(file_path):
    """
    Converts an Excel file to CSV format.
    """
    file_path = clean_path(file_path)
    csv_content = ""
    try:
        df = pd.read_excel(file_path)
        ouput = StringIO()
        df.to_csv(file_path, index=False)
        csv_content = output.getvalue()
    except Exception as e:
        print(f"Failed to convert Excel to CSV: {e}")

    return csv_content

def text_from_image(file_path):
    """
    Extracts plain text from an image using OCR.
    """
    file_path = clean_path(file_path)
    try:
        with Image.open(file_path) as img:
            # Perform OCR to extract text
            extracted_text = pytesseract.image_to_string(img).strip()
            return extracted_text if extracted_text else "[No text could be extracted from the image]"
    except Exception as e:
        print(f"Failed to process image: {file_path}, Error: {e}")
        return None

def text_from_other(file_path):
    """
    Extracts information from a file of unknown or unsupported type and returns plain text output.
    """
    file_path = clean_path(file_path)
    file_stats = os.stat(file_path)
    creation_time = datetime.fromtimestamp(file_stats.st_ctime)
    modified_time = datetime.fromtimestamp(file_stats.st_mtime)
    
    file_info = {
        "File Path": str(file_path),
        "File Size (bytes)": str(file_stats.st_size),
        "Creation Time": str(creation_time),
        "Modification Time": str(modified_time),
        "Permissions": oct(file_stats.st_mode & 0o777),
        "MIME Type": str(magic.from_file(file_path, mime=True)),
        "Hashes": {},
        "Readable Strings": [],
        "Magic Numbers": None,
        "Embedded URLs": [],
        "Entropy": None,
        "Exif Data": {},
    }

    # Get EXIF data
    exif_data = extract_exif(file_path)
    if exif_data:
        for key, value in exif_data.items():
            file_info["Exif Data"][key] = value

    def calculate_entropy(data):
        """Calculate Shannon entropy to assess randomness in the file."""
        if not data:
            return "0"
        counter = Counter(data)
        length = len(data)
        entropy = -sum((count / length) * math.log2(count / length) for count in counter.values())
        return str(entropy)

    def extract_strings(data):
        """Extract readable ASCII and Unicode strings."""
        ascii_regex = re.compile(rb'[ -~]{4,}')  # ASCII strings of length >= 4
        unicode_regex = re.compile(rb'(?:[\x20-\x7E][\x00]){4,}')  # Unicode UTF-16 strings
        strings = []
        strings.extend(match.decode('ascii') for match in ascii_regex.findall(data))
        strings.extend(match.decode('utf-16', errors='ignore') for match in unicode_regex.findall(data))
        return strings

    # Read the file as binary
    try:
        with open(file_path, 'rb') as file:
            binary_data = file.read()
            file_info["Hashes"]["SHA-256"] = hashlib.sha256(binary_data).hexdigest()
            file_info["Hashes"]["MD5"] = hashlib.md5(binary_data).hexdigest()
            file_info["Readable Strings"] = extract_strings(binary_data)[:10]  # Limit to 10 strings
            file_info["Entropy"] = calculate_entropy(binary_data)
            file_info["Magic Numbers"] = binary_data[:4].hex()
    except Exception as e:
        print(f"Error processing binary file {file_path}: {e}")
        return None

    # Generate plain text report
    report = [f"{key}: {value}" for key, value in file_info.items() if value]

    return "\n".join(report)

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract text from a file (PDF, DOCX, image, audio, etc.)")
    parser.add_argument("file", type=str, help="Path to the input file to extract text from")

    args = parser.parse_args()
    file_path = args.file

    result = extract_text(file_path)

    if result:
        print("\n--- Extracted Text ---\n")
        print(result)
    else:
        print("No text could be extracted or an error occurred.")

